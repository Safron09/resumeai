"""
Export service: generates PDF (WeasyPrint) and DOCX (python-docx),
optionally uploads both to Cloudinary if CLOUDINARY_URL is set.
"""
import io
import logging
import os
from pathlib import Path

import cloudinary.uploader
from docx import Document
from jinja2 import Environment, FileSystemLoader

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / 'templates'


# ── HTML / PDF ────────────────────────────────────────────────────────────────

def _render_html(content: dict, template_name: str = 'Classic') -> str:
    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    tpl = env.get_template('resume.html')
    return tpl.render(content=content, template_name=template_name)


def generate_pdf(content: dict, template_name: str = 'Classic') -> bytes:
    """Render resume HTML → PDF bytes via WeasyPrint."""
    from weasyprint import HTML  # lazy import keeps import time low for normal requests
    html = _render_html(content, template_name)
    return HTML(string=html).write_pdf()


# ── DOCX ──────────────────────────────────────────────────────────────────────

def generate_docx(content: dict) -> bytes:
    """Build a DOCX document from resume content dict."""
    doc = Document()

    # ── Name
    name = content.get('name', 'Resume')
    doc.add_heading(name, 0)

    # ── Contact line
    contact = content.get('contact', {})
    if isinstance(contact, dict):
        parts = [
            contact.get('email', ''),
            contact.get('phone', ''),
            contact.get('location', ''),
            contact.get('linkedin', ''),
        ]
        line = ' | '.join(p for p in parts if p)
        if line:
            doc.add_paragraph(line)

    # ── Summary
    if content.get('summary'):
        doc.add_heading('Summary', 1)
        doc.add_paragraph(content['summary'])

    # ── Experience
    experience = content.get('experience', [])
    if experience:
        doc.add_heading('Experience', 1)
        for job in experience:
            title = job.get('title', '')
            company = job.get('company', '')
            heading = f"{title} — {company}" if company else title
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
            if job.get('dates'):
                doc.add_paragraph(job['dates'])
            for bullet in job.get('bullets', []):
                doc.add_paragraph(f'\u2022 {bullet}')

    # ── Education
    education = content.get('education', [])
    if education:
        doc.add_heading('Education', 1)
        edu_list = education if isinstance(education, list) else [education]
        for edu in edu_list:
            if isinstance(edu, dict):
                degree = edu.get('degree', '')
                school = edu.get('school', '')
                dates = edu.get('dates', '')
                line = degree
                if school:
                    line += f' — {school}'
                if dates:
                    line += f' ({dates})'
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(str(edu))

    # ── Skills
    skills = content.get('skills', [])
    if skills:
        doc.add_heading('Skills', 1)
        if isinstance(skills, list):
            doc.add_paragraph(', '.join(str(s) for s in skills))
        else:
            doc.add_paragraph(str(skills))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Cloudinary upload ─────────────────────────────────────────────────────────

def _cloudinary_configured() -> bool:
    return bool(os.getenv('CLOUDINARY_URL', '').strip())


def _upload(data: bytes, public_id: str) -> str:
    result = cloudinary.uploader.upload(
        data,
        public_id=public_id,
        resource_type='raw',
        overwrite=True,
    )
    return result['secure_url']


# ── Main entry point ──────────────────────────────────────────────────────────

def export_resume(content: dict, job_id: int, template_name: str = 'Classic') -> tuple[str, str]:
    """
    Generate PDF + DOCX and (if Cloudinary is configured) upload both.
    Returns (pdf_url, docx_url) — empty strings if Cloudinary is not set up.
    """
    pdf_url = docx_url = ''

    try:
        pdf_bytes = generate_pdf(content, template_name)
        docx_bytes = generate_docx(content)

        if _cloudinary_configured():
            pdf_url = _upload(pdf_bytes, f'resumes/job_{job_id}.pdf')
            docx_url = _upload(docx_bytes, f'resumes/job_{job_id}.docx')
            logger.info('Uploaded PDF/DOCX to Cloudinary for job %s', job_id)
        else:
            logger.info('CLOUDINARY_URL not set — skipping upload for job %s', job_id)

    except Exception:
        logger.exception('export_resume failed for job %s', job_id)

    return pdf_url, docx_url
